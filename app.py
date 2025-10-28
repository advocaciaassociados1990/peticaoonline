import sys
import os

def resource_path(relative_path):
    """
    Retorna o caminho absoluto para um recurso, compatível com PyInstaller.
    """
    if getattr(sys, 'frozen', False):  # se executado como exe
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# === Estrutura de pastas ===
ROOT = resource_path("peticao_giovanna_simoes")
PASTA_BLOCOS = os.path.join(ROOT, "blocos")
PASTA_SAIDAS = os.path.join(ROOT, "saidas")
LOGO_CAMINHO = os.path.join(PASTA_BLOCOS, "logo_quebra_cabeca.png")

# Garante que as pastas existam mesmo em ambientes limitados (como Streamlit Cloud)
try:
    os.makedirs(PASTA_BLOCOS, exist_ok=True)
    os.makedirs(PASTA_SAIDAS, exist_ok=True)
except Exception:
    # fallback: cria uma pasta local no diretório atual, caso a estrutura acima não possa ser usada
    ROOT = os.getcwd()
    PASTA_BLOCOS = os.path.join(ROOT, "blocos")
    PASTA_SAIDAS = os.path.join(ROOT, "saidas")
    os.makedirs(PASTA_BLOCOS, exist_ok=True)
    os.makedirs(PASTA_SAIDAS, exist_ok=True)

import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageTk
import ttkbootstrap as tb

# === Função para ler bloco ===
def ler_bloco(nome_arquivo):
    caminho = os.path.join(PASTA_BLOCOS, nome_arquivo)
    if not os.path.exists(caminho):
        return f"⚠️ [Arquivo ausente: {nome_arquivo}]"
    with open(caminho, "r", encoding="utf-8") as f:
        return f.read().strip()

# === Processar marcação [LARANJA] ===
def processar_laranja(texto, paragrafo):
    partes = texto.split("[LARANJA]")
    for i, parte in enumerate(partes):
        if i % 2 == 0:
            paragrafo.add_run(parte)
        else:
            laranja, *resto = parte.split("[/LARANJA]")
            run = paragrafo.add_run(laranja)
            try:
                run.font.highlight_color = 6
            except Exception:
                pass
            if resto:
                paragrafo.add_run(resto[0])

# === Salvar DOCX ===
def salvar_peticao(texto_final, nome_arquivo="peticao_final.docx"):
    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11.5)
    try:
        estilo._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    for bloco in texto_final.split("\n\n"):
        if not bloco.strip():
            continue
        if "[PARAGRAFO]" in bloco:
            doc.add_paragraph()
            continue

        recuo_completo = "[RECUO_COMPLETO]" in bloco
        sem_recuo = "[SEM_RECUO]" in bloco
        centralizado = "[CENTRALIZADO]" in bloco
        bloco = (
            bloco.replace("[RECUO_COMPLETO]", "")
            .replace("[SEM_RECUO]", "")
            .replace("[CENTRALIZADO]", "")
        )

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if centralizado:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif recuo_completo:
            p.paragraph_format.left_indent = Cm(4)
            p.paragraph_format.first_line_indent = Cm(0)
        elif sem_recuo:
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            p.paragraph_format.first_line_indent = Cm(4)

        partes = bloco.split("[NEGRITO]")
        for i, parte in enumerate(partes):
            if i % 2 == 0:
                processar_laranja(parte, p)
            else:
                negrito, *resto = parte.split("[/NEGRITO]")
                run = p.add_run(negrito)
                run.bold = True
                if resto:
                    processar_laranja(resto[0], p)

    # === SALVAMENTO SEGURO ===
    try:
        os.makedirs(PASTA_SAIDAS, exist_ok=True)
        caminho_saida = os.path.join(PASTA_SAIDAS, nome_arquivo)
        doc.save(caminho_saida)
    except Exception:
        # fallback se o Streamlit Cloud não permitir gravar dentro da pasta
        caminho_saida = os.path.join(os.getcwd(), nome_arquivo)
        doc.save(caminho_saida)

    return caminho_saida

# === Montar texto ===
def montar_texto(dados):
    texto = ""
    texto += ler_bloco("bloco1_comarca.txt").replace("[INSERIR COMARCA]", dados["comarca"]) + "\n\n"
    texto += ler_bloco("bloco2_qualificacao_completa.txt").replace("[INSERIR QUALIFICAÇÃO INFORMADA]", dados["requerente"]) + "\n\n"
    texto += ler_bloco(f"bloco3_plano_{dados['plano']}.txt") + "\n\n"

    if dados["prioridade"] != "NENHUMA":
        bloco = ler_bloco(f"bloco4_prioridade_{dados['prioridade']}.txt")
        bloco = bloco.replace("[DESCREVER DOENÇA]", dados["doenca"])
        texto += bloco + "\n\n"

    texto += ler_bloco(f"bloco5_gratuidade_{dados['gratuidade']}.txt") + "\n\n"
    texto += ler_bloco("bloco6_fatos.txt").replace("[DESCREVER DOENÇA]", dados["doenca"]) + "\n\n"
    texto += ler_bloco(f"bloco7_negativa_{dados['negativa']}.txt") + "\n\n"
    texto += ler_bloco("bloco8_cdc.txt") + "\n\n"

    bloco9 = ler_bloco(f"bloco9_tipo_{dados['tipo_demanda']}.txt")
    bloco9 = bloco9.replace("[DESCREVER DOENÇA]", dados["doenca"])
    texto += bloco9 + "\n\n"

    texto += ler_bloco(f"bloco10_urgencia_{dados['urgencia_tipo']}.txt").replace("[DESCRIÇÃO URGÊNCIA]", dados["urgencia"]) + "\n\n"
    texto += ler_bloco(f"bloco11_pedidos_{dados['pedido']}.txt")

    return texto

# === Gerar petição ===
def gerar_peticao():
    dados = {
        "comarca": entrada_comarca.get().strip(),
        "requerente": entrada_requerente.get().strip(),
        "plano": combo_plano.get().strip().lower(),
        "prioridade": combo_prioridade.get().strip().upper(),
        "doenca": entrada_doenca.get().strip(),
        "negativa": combo_negativa.get().strip().lower(),
        "tipo_demanda": combo_tipo.get().strip().lower(),
        "pedido": combo_pedido.get().strip().lower(),
        "urgencia": entrada_urgencia.get("1.0", "end-1c").strip(),
        "urgencia_tipo": combo_urgencia_tipo.get().strip().lower(),
        "gratuidade": combo_gratuidade.get().strip().upper() or "NENHUMA",
    }

    campos_obrigatorios = ["comarca", "requerente", "plano", "doenca", "negativa", "tipo_demanda", "pedido"]
    for c in campos_obrigatorios:
        if not dados[c]:
            messagebox.showwarning("ATENÇÃO", "PREENCHA TODOS OS CAMPOS OBRIGATÓRIOS ANTES DE GERAR A PETIÇÃO.")
            return

    texto = montar_texto(dados)
    nome_arquivo = f"Peticao_{dados['requerente'].split()[0]}_{dados['comarca'].replace(' ', '_')}.docx"
    caminho_saida = salvar_peticao(texto, nome_arquivo)
    messagebox.showinfo("PETIÇÃO GERADA!", f"ARQUIVO SALVO EM:\n{caminho_saida}")

# === INTERFACE ===
app = tb.Window(themename="flatly")
app.title("🧩 GERADOR DE PETIÇÕES")
app.geometry("680x820")

cabecalho_text = (
    "Desenvolvido por Ana Paula Braga para Giovanna Rocha Simões — Compartilhamento proibido.\n"
    "Modelo básico, a ser refinado de acordo com o caso específico."
)
tk.Label(app, text=cabecalho_text, font=("Calibri", 10), justify="center", wraplength=620).pack(pady=(8, 6))

cores = ["#4C9ED9", "#F5D142", "#E94E77", "#53C28B"]
for i in range(8):
    frame = tk.Frame(app, bg=cores[i % len(cores)], height=8)
    frame.pack(fill="x", pady=(0 if i == 0 else 1))

if os.path.exists(LOGO_CAMINHO):
    try:
        img = Image.open(LOGO_CAMINHO).resize((80, 80))
        logo_img = ImageTk.PhotoImage(img)
        tk.Label(app, image=logo_img).pack(pady=8)
    except Exception:
        pass

container = ttk.Frame(app)
container.pack(fill="both", expand=True)
canvas = tk.Canvas(container)
scrollbar = ttk.Scrollbar(container, orient="vertical", command=canvas.yview)
scrollable_frame = ttk.Frame(canvas)
scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
canvas.configure(yscrollcommand=scrollbar.set)
canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")
frame_principal = scrollable_frame

# === CAMPOS ===
ttk.Label(frame_principal, text="COMARCA:").pack(anchor="w", pady=(6, 0))
entrada_comarca = ttk.Entry(frame_principal)
entrada_comarca.pack(fill="x")

ttk.Label(frame_principal, text="REQUERENTE:").pack(anchor="w", pady=(6, 0))
entrada_requerente = ttk.Entry(frame_principal)
entrada_requerente.pack(fill="x")

ttk.Label(frame_principal, text="PLANO DE SAÚDE:").pack(anchor="w", pady=(6, 0))
combo_plano = ttk.Combobox(
    frame_principal,
    values=["UNIMED", "BRADESCO", "NOTREDAME", "SAMARITANO", "AMIL", "SULAMERICA"],
    state="readonly"
)
combo_plano.pack(fill="x")

ttk.Label(frame_principal, text="PRIORIDADE DE TRAMITAÇÃO:").pack(anchor="w", pady=(6, 0))
combo_prioridade = ttk.Combobox(
    frame_principal,
    values=["NENHUMA", "IDOSO", "DEFICIENTE"],
    state="readonly"
)
combo_prioridade.pack(fill="x")

ttk.Label(frame_principal, text="GRATUIDADE DE JUSTIÇA:").pack(anchor="w", pady=(6, 0))
combo_gratuidade = ttk.Combobox(
    frame_principal,
    values=["NENHUMA", "IDOSO_OU_TUTELADO", "MENOR"],
    state="readonly"
)
combo_gratuidade.pack(fill="x")

ttk.Label(frame_principal, text="DOENÇA / CONDIÇÃO:").pack(anchor="w", pady=(6, 0))
entrada_doenca = ttk.Entry(frame_principal)
entrada_doenca.pack(fill="x")

ttk.Label(frame_principal, text="TIPO DE NEGATIVA:").pack(anchor="w", pady=(6, 0))
combo_negativa = ttk.Combobox(frame_principal, values=["TACITA", "OUTRA"], state="readonly")
combo_negativa.pack(fill="x")

ttk.Label(frame_principal, text="TIPO DE DEMANDA:").pack(anchor="w", pady=(6, 0))
combo_tipo = ttk.Combobox(
    frame_principal,
    values=[
        "deficiencia_clinico",
        "deficiencia_domiciliar",
        "idoso_clinico",
        "idoso_domiciliar",
        "outros"
    ],
    state="readonly"
)
combo_tipo.pack(fill="x")

ttk.Label(frame_principal, text="URGÊNCIA (DESCRIÇÃO):").pack(anchor="w", pady=(6, 0))
entrada_urgencia = tk.Text(frame_principal, height=4)
entrada_urgencia.pack(fill="x")

ttk.Label(frame_principal, text="TIPO DE URGÊNCIA:").pack(anchor="w", pady=(6, 0))
combo_urgencia_tipo = ttk.Combobox(
    frame_principal,
    values=["clinica", "domiciliar"],
    state="readonly"
)
combo_urgencia_tipo.pack(fill="x")

ttk.Label(frame_principal, text="PEDIDO:").pack(anchor="w", pady=(6, 0))
combo_pedido = ttk.Combobox(
    frame_principal,
    values=["clinica", "domiciliar"],
    state="readonly"
)
combo_pedido.pack(fill="x")

ttk.Button(frame_principal, text="🧩 GERAR PETIÇÃO", command=gerar_peticao).pack(pady=12)

app.mainloop()
